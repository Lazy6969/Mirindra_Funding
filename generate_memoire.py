from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_memoire():
    doc = Document()
    
    # --- Page de Garde ---
    title = doc.add_heading('MÉMOIRE DE FIN D\'ÉTUDES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\nSujet : Conception et Réalisation d\'une Plateforme de Crowdfunding pour Projets Éducatifs\n"MIRINDRA FUNDING"')
    run.bold = True
    run.font.size = Pt(18)

    doc.add_page_break()

    # --- Sommaire ---
    doc.add_heading('Sommaire', level=1)
    sections = [
        "Introduction Générale",
        "PARTIE I : ANALYSE ET ÉTUDE PRÉALABLE (MÉTHODE MERISE)",
        "  Chapitre 1 : Présentation du projet et problématique",
        "  Chapitre 2 : Analyse fonctionnelle et Modèle Conceptuel de Données (MCD)",
        "  Chapitre 3 : Modèle Conceptuel des Traitements (MCT)",
        "PARTIE II : CONCEPTION TECHNIQUE ET ARCHITECTURE",
        "  Chapitre 4 : Passage au Modèle Logique de Données (MLD)",
        "  Chapitre 5 : Architecture logicielle (PHP/PDO/MVC)",
        "  Chapitre 6 : Design et Expérience Utilisateur (UX/UI)",
        "PARTIE III : RÉALISATION ET DÉPLOIEMENT",
        "  Chapitre 7 : Implémentation de la base de données et du Backend",
        "  Chapitre 8 : Sécurité et Validation des Projets",
        "  Chapitre 9 : Manuel Utilisateur et Administrateur",
        "Conclusion Générale",
        "Bibliographie",
        "Annexes"
    ]
    for section in sections:
        doc.add_paragraph(section)

    doc.add_page_break()

    # --- Introduction ---
    doc.add_heading('Introduction Générale', level=1)
    doc.add_paragraph(
        "À l'ère du numérique, le financement participatif ou crowdfunding s'impose comme un levier majeur "
        "pour l'innovation sociale. Mirindra Funding naît de la nécessité de soutenir le secteur éducatif "
        "à Madagascar en permettant aux enseignants et élèves de lever des fonds pour des projets concrets."
    )

    # --- Merise Partie I ---
    doc.add_heading('PARTIE I : ANALYSE ET ÉTUDE PRÉALABLE', level=1)
    doc.add_heading('Le Modèle Conceptuel de Données (MCD)', level=2)
    doc.add_paragraph(
        "Dans le cadre de Mirindra Funding, nous avons identifié les entités suivantes :\n"
        "- UTILISATEUR (Email, Password, Type, École)\n"
        "- PROJET (Titre, Description, Objectif, Statut)\n"
        "- DON (Montant, Date, Anonymat)\n"
        "- CATÉGORIE (Nom, Couleur)"
    )
    doc.add_paragraph("Relations : Un Utilisateur 'Crée' un ou plusieurs Projets. Un Utilisateur 'Effectue' un Don pour un Projet.")

    # --- Merise Partie II ---
    doc.add_heading('PARTIE II : CONCEPTION TECHNIQUE', level=1)
    doc.add_heading('Le Modèle Logique de Données (MLD)', level=2)
    doc.add_paragraph(
        "Le passage du MCD au MLD donne la structure SQL suivante :\n"
        "- users (id, email, password, first_name, last_name, user_type...)\n"
        "- projects (id, #user_id, title, target_amount, status...)\n"
        "- donations (id, #project_id, #user_id, amount, donor_name...)"
    )

    # --- Partie III Implementation ---
    doc.add_heading('PARTIE III : RÉALISATION', level=1)
    doc.add_paragraph(
        "L'application repose sur une stack LAMP (Linux, Apache, MySQL, PHP). "
        "L'interactivité est renforcée par l'usage de JavaScript vanilla pour le filtrage "
        "en temps réel des projets et des effets visuels de curseur."
    )

    # --- Conclusion ---
    doc.add_heading('Conclusion Générale', level=1)
    doc.add_paragraph(
        "Le projet Mirindra Funding a permis de mettre en place une solution technique viable "
        "et sécurisée pour le financement de l'éducation. Les perspectives d'évolution incluent "
        "l'intégration de paiements mobiles (Mobile Money) locaux."
    )

    # --- Bibliographie ---
    doc.add_heading('Bibliographie', level=1)
    doc.add_paragraph("1. Hubert Tardieu, La méthode Merise, Tome 1 : Principes et outils.")
    doc.add_paragraph("2. Documentation PHP officielle (php.net).")
    doc.add_paragraph("3. MDN Web Docs pour les animations CSS et JS.")

    # --- Annexes ---
    doc.add_heading('Annexes', level=1)
    doc.add_heading('Schéma de la base de données (SQL)', level=2)
    sql_code = """
    CREATE TABLE projects (
        id INT PRIMARY KEY AUTO_INCREMENT,
        user_id INT,
        title VARCHAR(255),
        target_amount DECIMAL(10,2),
        status ENUM('pending', 'active', 'completed')
    );
    """
    doc.add_paragraph(sql_code)

    # Sauvegarde
    doc.save('Memoire_Mirindra_Funding.docx')
    print("Fichier 'Memoire_Mirindra_Funding.docx' généré avec succès.")

if __name__ == "__main__":
    try:
        create_memoire()
    except ImportError:
        print("Veuillez installer python-docx : pip install python-docx")